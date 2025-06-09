import pandas as pd
import matplotlib.pyplot as plt
from datetime import datetime
from docx import Document
from docx.shared import Inches

def generate_report():
    # Initialize a Word document
    doc = Document()
    doc.add_heading('Interim Report - Week 1 Challenge', level=1)

    # Section 1: Progress Summary
    doc.add_heading('Progress Summary', level=2)
    doc.add_paragraph(
        "This interim report covers the progress made on Task 1 (Git and GitHub) and partial progress on Task 2 (Quantitative Analysis). "
        "Below are the key findings and methodologies employed so far."
    )

    # Section 2: Task 1 - Git and GitHub
    doc.add_heading('Task 1: Git and GitHub', level=2)
    doc.add_paragraph(
        "- Repository setup completed with the required folder structure.\n"
        "- Exploratory Data Analysis (EDA) performed on the dataset, including:\n"
        "  - Descriptive statistics for headline lengths.\n"
        "  - Analysis of publication frequency over time.\n"
        "  - Publisher contribution analysis."
    )

    # Add a placeholder plot (replace with actual data)
    plt.figure(figsize=(6, 4))
    plt.plot([1, 2, 3], [4, 5, 6])
    plt.title('Sample Plot: Publication Frequency Over Time')
    plt.savefig('publication_frequency.png')
    doc.add_picture('publication_frequency.png', width=Inches(5))

    # Section 3: Task 2 - Quantitative Analysis
    doc.add_heading('Task 2: Quantitative Analysis', level=2)
    doc.add_paragraph(
        "- Stock price data loaded and prepared for analysis.\n"
        "- Technical indicators (e.g., Moving Averages, RSI) calculated using TA-Lib.\n"
        "- Preliminary visualizations created to explore trends."
    )

    # Section 4: Challenges
    doc.add_heading('Challenges Encountered', level=2)
    doc.add_paragraph(
        "- Data alignment issues between news and stock datasets.\n"
        "- Sentiment analysis accuracy needs further refinement."
    )

    # Save the report
    doc.save('reports/interim_report.docx')

    # Convert to PDF (requires pandoc and LaTeX)
    import subprocess
    subprocess.run(['pandoc', 'reports/interim_report.docx', '-o', 'reports/interim_report.pdf'])

if __name__ == '__main__':
    generate_report()